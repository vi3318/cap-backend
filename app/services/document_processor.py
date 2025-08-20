"""
Document Processing Service
Handles document upload, OCR, text extraction, and preprocessing for legal documents
"""

import os
import logging
from typing import Dict, Any, List, Optional
from pathlib import Path
import PyPDF2
import pdfplumber
from PIL import Image
import pytesseract
from langdetect import detect, LangDetectException
import io
import mimetypes
import hashlib
import time
from datetime import datetime

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self, settings):
        self.settings = settings
        self.upload_dir = Path(settings.upload_dir)
        self.upload_dir.mkdir(exist_ok=True)
        
        # Supported file types
        self.supported_types = {
            'application/pdf': self._process_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._process_docx,
            'text/plain': self._process_text,
            'image/jpeg': self._process_image,
            'image/png': self._process_image,
            'image/tiff': self._process_image,
            'image/bmp': self._process_image
        }
    
    async def process_document(self, file_content: bytes, filename: str, content_type: str) -> Dict[str, Any]:
        """
        Process uploaded document and extract text content.
        """
        try:
            # Generate document ID
            document_id = self._generate_document_id(file_content, filename)
            
            # Save file
            file_path = self.upload_dir / f"{document_id}_{filename}"
            with open(file_path, 'wb') as f:
                f.write(file_content)
            
            # Process based on content type
            if content_type in self.supported_types:
                processor = self.supported_types[content_type]
                text_content = await processor(file_path)
            else:
                # Try to determine type from file extension
                ext = Path(filename).suffix.lower()
                if ext == '.pdf':
                    text_content = await self._process_pdf(file_path)
                elif ext == '.docx':
                    text_content = await self._process_docx(file_path)
                elif ext in ['.txt', '.md']:
                    text_content = await self._process_text(file_path)
                elif ext in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
                    text_content = await self._process_image(file_path)
                else:
                    raise ValueError(f"Unsupported file type: {content_type}")
            
            # Detect language
            language = self._detect_language(text_content)
            
            # Extract basic entities (simplified)
            entities = self._extract_basic_entities(text_content)
            
            return {
                'document_id': document_id,
                'filename': filename,
                'file_path': str(file_path),
                'content_type': content_type,
                'file_size': len(file_content),
                'text_content': text_content,
                'language': language,
                'entities': entities,
                'upload_time': datetime.now().isoformat(),
                'status': 'processed'
            }
            
        except Exception as e:
            logger.error(f"Error processing document {filename}: {e}")
            raise
    
    async def _process_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        try:
            text_content = ""
            
            # Try pdfplumber first (better for complex PDFs)
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += page_text + "\n"
            except Exception as e:
                logger.warning(f"pdfplumber failed, trying PyPDF2: {e}")
                
                # Fallback to PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += page_text + "\n"
            
            return text_content.strip()
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            raise
    
    async def _process_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            # Import here to avoid dependency issues
            from docx import Document
            
            doc = Document(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + "\n"
            
            return text_content.strip()
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            raise
    
    async def _process_text(self, file_path: Path) -> str:
        """Extract text from plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read().strip()
                except UnicodeDecodeError:
                    continue
            raise ValueError("Could not decode text file with any encoding")
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {e}")
            raise
    
    async def _process_image(self, file_path: Path) -> str:
        """Extract text from image using OCR."""
        try:
            # Open image
            image = Image.open(file_path)
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Extract text using OCR
            text_content = pytesseract.image_to_string(image)
            
            return text_content.strip()
            
        except Exception as e:
            logger.error(f"Error processing image {file_path}: {e}")
            raise
    
    def _detect_language(self, text: str) -> str:
        """Detect language of text content."""
        try:
            if not text or len(text.strip()) < 10:
                return 'unknown'
            
            # Use first 1000 characters for language detection
            sample_text = text[:1000]
            language = detect(sample_text)
            return language
            
        except LangDetectException:
            return 'unknown'
        except Exception as e:
            logger.warning(f"Language detection failed: {e}")
            return 'unknown'
    
    def _extract_basic_entities(self, text: str) -> List[Dict[str, str]]:
        """
        Extract basic entities from text (simplified version).
        This is a basic implementation without heavy ML dependencies.
        """
        entities = []
        
        if not text:
            return entities
        
        # Simple pattern matching for common entities
        import re
        
        # Extract potential dates
        date_patterns = [
            r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b',  # DD/MM/YYYY
            r'\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b',    # YYYY/MM/DD
            r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{1,2},? \d{4}\b'  # Month DD, YYYY
        ]
        
        for pattern in date_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                entities.append({
                    'text': match.group(),
                    'type': 'DATE',
                    'start': match.start(),
                    'end': match.end()
                })
        
        # Extract potential organizations (words starting with capital letters)
        org_pattern = r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b'
        matches = re.finditer(org_pattern, text)
        for match in matches:
            # Filter out common words that aren't organizations
            word = match.group()
            if len(word.split()) >= 2 and word not in ['The', 'And', 'For', 'With', 'From', 'This', 'That']:
                entities.append({
                    'text': word,
                    'type': 'ORGANIZATION',
                    'start': match.start(),
                    'end': match.end()
                })
        
        # Extract potential monetary amounts
        money_pattern = r'\$\d+(?:,\d{3})*(?:\.\d{2})?|\d+(?:,\d{3})*(?:\.\d{2})?\s*(?:dollars?|USD|EUR|GBP)'
        matches = re.finditer(money_pattern, text, re.IGNORECASE)
        for match in matches:
            entities.append({
                'text': match.group(),
                'type': 'MONEY',
                'start': match.start(),
                'end': match.end()
            })
        
        return entities[:20]  # Limit to first 20 entities
    
    def _generate_document_id(self, content: bytes, filename: str) -> str:
        """Generate unique document ID."""
        timestamp = str(int(time.time()))
        content_hash = hashlib.md5(content).hexdigest()[:8]
        filename_hash = hashlib.md5(filename.encode()).hexdigest()[:8]
        
        return f"{timestamp}_{content_hash}_{filename_hash}"
    
    async def get_document_info(self, document_id: str) -> Optional[Dict[str, Any]]:
        """Get information about a processed document."""
        try:
            # Look for files with this document ID
            for file_path in self.upload_dir.glob(f"{document_id}_*"):
                if file_path.exists():
                    stat = file_path.stat()
                    return {
                        'document_id': document_id,
                        'filename': file_path.name.replace(f"{document_id}_", ""),
                        'file_path': str(file_path),
                        'file_size': stat.st_size,
                        'upload_time': datetime.fromtimestamp(stat.st_mtime).isoformat(),
                        'exists': True
                    }
            return None
            
        except Exception as e:
            logger.error(f"Error getting document info for {document_id}: {e}")
            return None 